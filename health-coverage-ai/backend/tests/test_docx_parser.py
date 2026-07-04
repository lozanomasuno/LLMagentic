from pathlib import Path

from docx import Document

from app.ai.parsers.docx_parser import DocxParser


def test_docx_parser_produces_chunks(tmp_path: Path):
    doc_path = tmp_path / "DOC1_Test.docx"
    doc = Document()
    doc.add_paragraph("1. Objeto y alcance")
    doc.add_paragraph("Este documento define las condiciones generales del plan.")
    doc.add_paragraph("2. Servicios cubiertos")
    doc.add_paragraph("Incluye consulta externa y apoyo diagnóstico básico.")
    doc.save(str(doc_path))

    parser = DocxParser()
    chunks = parser.parse(doc_path)

    assert len(chunks) >= 1
    assert chunks[0].document_slug.startswith("doc1")
    assert all(chunk.char_count > 0 for chunk in chunks)
